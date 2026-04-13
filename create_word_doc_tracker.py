import os
import sys

try:
    import docx
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    import docx

doc = docx.Document()
doc.add_heading('Deployment Progress & Strategy Tracker', 0)

# Section 1
doc.add_heading('1. What We Have Done So Far', level=2)

doc.add_heading('Step 1: Database (Supabase) Setup', level=3)
doc.add_paragraph("What it is: Supabase is serving as our remote SQL database.", style='List Bullet')
doc.add_paragraph("Why we did it: A live web application cannot read a local .csv file on your laptop. We needed a live data source in the cloud.", style='List Bullet')
doc.add_paragraph("Status: DONE.", style='List Bullet')

doc.add_heading('Step 2: Frontend (Dashboard) Deployment', level=3)
doc.add_paragraph("What it is: The visual interface where users can see inventory alerts and predictions.", style='List Bullet')
doc.add_paragraph("Why we did it: So stakeholders can access the predictions anywhere via a web link.", style='List Bullet')
doc.add_paragraph("Status: DONE. The frontend is deployed and linked.", style='List Bullet')

# Section 2
doc.add_heading('2. What is Next? (The "Missing Link")', level=2)
doc.add_paragraph("Right now, you have live Data (Supabase) and a Visual Dashboard (Frontend). But the dashboard isn't getting daily \"smart predictions\" from your AI Model.")
doc.add_paragraph("You need a \"Brain\" (a cloud computer) that wakes up every night, grabs the newest data from Supabase, runs it through your XGBoost AI model, and sends the updated predictions to your Dashboard.")

doc.add_heading('Phase 3: Setup GitHub Actions (Backend Automation)', level=3)
doc.add_paragraph("What it is: GitHub Actions is a free automation tool. It acts as our invisible \"cloud computer.\"", style='List Bullet')
doc.add_paragraph("Why we need it: We cannot expect you to manually run your script on your own laptop every night. GitHub Actions will automate this process completely for free.", style='List Bullet')

# Section 3
doc.add_heading('3. How to Execute Phase 3 (The Final Step)', level=2)
doc.add_paragraph("Here is the exact step-by-step process we need to follow to finish the deployment:")

doc.add_heading('Step 3.1: Connect Python to Supabase', level=3)
doc.add_paragraph("Action Needed: Your script likely still uses pd.read_csv() to read local files. We need to update that specific part of the code so that it connects to your live Supabase database url using psycopg2 or SQLAlchemy to read the live table instead.", style='List Bullet')

doc.add_heading('Step 3.2: Get Your Code on GitHub', level=3)
doc.add_paragraph("Action Needed: If your project isn't already there, we need to push all your files (Python scripts, best_model.pkl, etc.) to a GitHub repository online, so GitHub has access to the AI model.", style='List Bullet')

doc.add_heading('Step 3.3: Write the Automation "CRON" Job', level=3)
doc.add_paragraph("Action Needed: We will create a folder structure named .github/workflows/ in your project.", style='List Bullet')
doc.add_paragraph("We will place a file called daily_forecast.yml inside it. In that file, we write an instruction that tells GitHub: \"Every night at 1:00 AM, turn on a server, run the Python program, generate the JSON output, and save it for the website to read.\"", style='List Bullet')
doc.add_paragraph("Once this is done, your architecture is 100% complete! The data will legally flow from Supabase -> GitHub (Processing) -> Frontend.")

doc.save('Deployment_Tracker.docx')
print("Successfully created Deployment_Tracker.docx")
