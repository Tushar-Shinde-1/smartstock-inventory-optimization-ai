"""
SmartStock — Create Deployment & Architecture Documentation (DOCX)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ---- TITLE ----
t = doc.add_heading('CRISP-ML(Q) Phase 6: Deployment & Integration Strategy', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('SmartStock — Real-Time Automated Data Pipeline')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(31,78,121)
doc.add_paragraph()

# ---- 1. Architecture Overview ----
doc.add_heading('1. Architecture Overview (Zero-Cost Stack)', level=1)
p = doc.add_paragraph('To transition from a static CSV file to an automated Cloud Architecture, we use three free tiers:')
bp(doc, 'Database Layer (Supabase): Free PostgreSQL database to store continuous daily transactions.')
bp(doc, 'Compute Layer (GitHub Actions): CRON jobs to automatically run inference scripts every night.')
bp(doc, 'Dashboard Layer (Vercel): Free tier CDN to host the interactive HTML/JS dashboard.')

doc.add_paragraph()

# ---- 2. Automated Pipeline Workflow ----
doc.add_heading('2. Daily Automation Workflow & Data Ingestion', level=1)
p = doc.add_paragraph('A Nightly Batch Processing model is mathematically superior for this inventory system, as the model relies on historical lag features (e.g., 7-day lags) that require complete daily aggregations.')

doc.add_heading('Phase 1: Continuous Data Collection', level=2)
bp(doc, 'Retail POS systems send real-time sales throughout the business hours.')
bp(doc, 'These HTTP POST requests insert transaction records immediately into the Supabase database.')

doc.add_heading('Phase 2: Nightly Inference (1:00 AM)', level=2)
bp(doc, 'GitHub Action wakes up on a CRON schedule (0 1 * * *).')
bp(doc, 'Pulls the most recent 30-days of completed sales from Supabase.')
bp(doc, 'Runs the pre-trained XGBoost model (best_model.pkl).')
bp(doc, 'Generates smartstock_insights.json with tomorrows predictions and alerts.')
bp(doc, 'Pushes the new JSON file back into the code repository.')

doc.add_heading('Phase 3: Automated Dashboard Refresh (Morning)', level=2)
bp(doc, 'The code push automatically triggers a Vercel rebuild.')
bp(doc, 'The visual dashboard is updated with exactly what needs to be ordered for the day.')

doc.add_paragraph()

# ---- 3. Managing Model Drift ----
doc.add_heading('3. Automated Monthly Retraining (Model Drift handling)', level=1)
p = doc.add_paragraph('To ensure the XGBoost model remains accurate as customer buying patterns shift:')
bp(doc, 'A separate GitHub Action (Monthly Retraining) triggers on the 1st of every month.')
bp(doc, 'Fetches all historical data and runs run_modeling.py.')
bp(doc, 'Tests the new model\'s R² and MAE.')
bp(doc, 'Overwrites best_model.pkl in production if accuracy holds.')

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Document Version: 1.0  |  Date: April 2025  |  SmartStock Deployment Architecture')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(128,128,128)

doc.save('06_Deployment_Strategy.docx')
print("✅ Created: 06_Deployment_Strategy.docx")
