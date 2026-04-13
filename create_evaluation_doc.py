"""
SmartStock — Create Phase 5 Evaluation & SmartStock Documentation (DOCX)
"""
import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load insights
with open('dashboard/smartstock_insights.json', 'r') as f:
    insights = json.load(f)

with open('dashboard/model_results.json', 'r') as f:
    model_data = json.load(f)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, hdr_color='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(cell, hdr_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 1: set_cell_shading(cell, 'F2F2F2')

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ---- TITLE ----
t = doc.add_heading('CRISP-ML(Q) Phase 5: Model Evaluation & SmartStock Insights', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('SmartStock — Intelligent Inventory Optimization Platform\nEvaluation, Business Intelligence & Recommendations Report')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(31,78,121)
doc.add_paragraph()

# ---- 1. EXECUTIVE SUMMARY ----
doc.add_heading('1. Executive Summary', level=1)
best = insights['evaluation']['best_model']
metrics = insights['evaluation']['test_metrics']
p = doc.add_paragraph()
p.add_run('Best Model: ').bold = True
p.add_run(f'{best} — selected based on highest test R² score.')
doc.add_paragraph()
add_table(doc, ['Metric', 'Value', 'Interpretation'], [
    ['MAE', str(metrics['MAE']), f'Average prediction error is {metrics["MAE"]:.1f} units'],
    ['RMSE', str(metrics['RMSE']), f'Root mean squared error penalizing large errors'],
    ['R²', str(metrics['R2']), f'Model explains {metrics["R2"]*100:.2f}% of demand variance'],
    ['MAPE', f'{metrics["MAPE"]}%', f'Average {metrics["MAPE"]:.1f}% error relative to actual demand'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Business Outcomes: ').bold = True
bp(doc, f'{len(insights["product_forecast"])} product categories forecasted')
bp(doc, f'{len(insights["inventory_optimization"])} product-brand combinations with Safety Stock & Reorder Points')
bp(doc, f'Stock health analyzed: {insights["health_summary"]}')
bp(doc, f'Daily recommendations generated: {insights["alert_summary"]}')

doc.add_paragraph()

# ---- 2. MODEL EVALUATION ----
doc.add_heading('2. Model Evaluation', level=1)

doc.add_heading('2.1 Model Comparison Summary', level=2)
comparison = insights['evaluation']['model_comparison']
rows = []
for name, m in comparison.items():
    rows.append([name, str(m['MAE']), str(m['RMSE']), str(m['R2']), f'{m["MAPE"]}%'])
add_table(doc, ['Model', 'MAE', 'RMSE', 'R²', 'MAPE'], rows)

doc.add_paragraph()
doc.add_heading('2.2 Cross-Validation Results', level=2)
p = doc.add_paragraph()
p.add_run('5-Fold CV provides a robust estimate of model generalization: ')
cv_rows = []
for name, cv in insights['evaluation']['cv_results'].items():
    cv_rows.append([name, f'{cv["mae_mean"]} ± {cv["mae_std"]}', f'{cv["r2_mean"]} ± {cv["r2_std"]}'])
add_table(doc, ['Model', 'CV MAE (mean ± std)', 'CV R² (mean ± std)'], cv_rows)

doc.add_paragraph()
doc.add_heading('2.3 Evaluation Criteria', level=2)
for criterion, desc in insights['evaluation']['evaluation_criteria'].items():
    bp(doc, f'{criterion.title()}: {desc}')

doc.add_paragraph()

# ---- 3. DEMAND FORECASTING ----
doc.add_heading('3. Demand Forecasting Results', level=1)
p = doc.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('Forecast product demand accurately using predictive models to enable proactive inventory management.')
doc.add_paragraph()

doc.add_heading('3.1 Product-Level Demand Forecast', level=2)
pf_rows = []
for pf in insights['product_forecast']:
    pf_rows.append([pf['product'], str(pf['avg_daily_demand']), str(pf['predicted_avg_demand']),
                     str(pf['total_predicted_30d']), str(pf['demand_std'])])
add_table(doc, ['Product', 'Historical Avg', 'Predicted Avg', '30-day Total', 'Std Dev'], pf_rows)

doc.add_paragraph()
doc.add_heading('3.2 Product × Brand Breakdown', level=2)
p = doc.add_paragraph('Top 15 product-brand demand forecasts:')
pbf_rows = []
for pbf in insights['product_brand_forecast'][:15]:
    pbf_rows.append([pbf['product'], pbf['brand'], str(pbf['avg_daily_demand']),
                      str(pbf['predicted_demand']), f'₹{pbf["avg_price"]}', str(pbf['current_stock'])])
add_table(doc, ['Product', 'Brand', 'Avg Daily', 'Predicted', 'Avg Price', 'Current Stock'], pbf_rows)

doc.add_paragraph()

# ---- 4. SAFETY STOCK & REORDER POINTS ----
doc.add_heading('4. Safety Stock & Reorder Point Optimization', level=1)
p = doc.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('Dynamically calculate optimal safety stock and reorder points to minimize stockouts while controlling holding costs.')

doc.add_paragraph()
doc.add_heading('4.1 Formulas Used', level=2)
p = doc.add_paragraph()
p.add_run('Safety Stock = Z × σ_demand × √(Lead Time)').italic = True
doc.add_paragraph()
bp(doc, 'Z = 1.65 (for 95% service level — 95% probability of not facing a stockout)')
bp(doc, 'σ_demand = standard deviation of historical demand')
bp(doc, 'Lead Time = 7 days (assumed supplier lead time)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Reorder Point = (Average Daily Demand × Lead Time) + Safety Stock').italic = True
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Economic Order Quantity (EOQ) = √(2 × D × S / H)').italic = True
doc.add_paragraph()
bp(doc, 'D = Annual demand (daily avg × 365)')
bp(doc, 'S = Fixed ordering cost (₹500 assumed)')
bp(doc, 'H = Holding cost per unit (20% of unit price)')

doc.add_paragraph()
doc.add_heading('4.2 Inventory Optimization Results', level=2)
inv_rows = []
for inv in insights['inventory_optimization'][:15]:
    inv_rows.append([inv['product'], inv['brand'], str(inv['avg_daily_demand']),
                      str(inv['safety_stock']), str(inv['reorder_point']),
                      str(inv['eoq']), str(inv['current_stock']), inv['stock_status']])
add_table(doc, ['Product', 'Brand', 'Avg Demand', 'Safety Stock', 'ROP', 'EOQ', 'Stock', 'Status'], inv_rows)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Stock Status Legend: ').bold = True
bp(doc, 'Critical — Current stock below 50% of reorder point (immediate action required)')
bp(doc, 'Reorder — Current stock below reorder point (place order)')
bp(doc, 'Healthy — Stock between ROP and 3× ROP (normal operations)')
bp(doc, 'Overstock — Stock exceeds 3× ROP (consider discounting)')

doc.add_paragraph()

# ---- 5. SLOW-MOVING & DEADSTOCK ----
doc.add_heading('5. Slow-Moving & Deadstock Identification', level=1)
p = doc.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('Identify slow-moving and deadstock items to reduce holding costs and free up working capital.')

doc.add_paragraph()
doc.add_heading('5.1 Classification Criteria', level=2)
add_table(doc, ['Category', 'Definition', 'Action'], [
    ['Fast-Moving', 'Turnover ratio > 8', 'Maintain consistent stock levels'],
    ['Normal', 'Average sales performance', 'Standard replenishment cycle'],
    ['Slow-Moving', 'Sales below 25th percentile', 'Run promotional discounts'],
    ['Deadstock', 'Sales below 10th percentile', 'Deep discount or bundle clearance'],
])

doc.add_paragraph()
doc.add_heading('5.2 Health Summary', level=2)
h_rows = [[k, str(v)] for k, v in insights['health_summary'].items()]
add_table(doc, ['Category', 'Product-Brand Count'], h_rows)

doc.add_paragraph()
doc.add_heading('5.3 Key Metrics', level=2)
p = doc.add_paragraph()
p.add_run('Inventory Turnover Ratio: ').bold = True
p.add_run('Total Quantity Sold / Average Stock. Higher values indicate faster-moving items.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Days of Supply: ').bold = True
p.add_run('Average Stock / Average Daily Sales. Lower values indicate faster stock depletion.')

doc.add_paragraph()

# ---- 6. DAILY RECOMMENDATIONS ----
doc.add_heading('6. Automated Daily Recommendations', level=1)
p = doc.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('Generate automated, actionable inventory recommendations to streamline daily operations.')

doc.add_paragraph()
doc.add_heading('6.1 Recommendation Logic', level=2)
add_table(doc, ['Alert Type', 'Trigger Condition', 'Recommended Action'], [
    ['Critical', 'Stock ≤ Safety Stock', 'Order EOQ units immediately (express shipping)'],
    ['Reorder Now', 'Stock ≤ Reorder Point', 'Place standard order for EOQ units'],
    ['Healthy', 'Stock between ROP and 3×ROP', 'No action — monitor normally'],
    ['Overstock Warning', 'Stock ≥ 3×ROP & low turnover', 'Consider 10-15% promotional discount'],
    ['Deadstock', 'Very low or zero recent sales', 'Clearance sale or bundle with fast-movers'],
])

doc.add_paragraph()
doc.add_heading('6.2 Alert Summary', level=2)
a_rows = [[k, str(v)] for k, v in insights['alert_summary'].items()]
add_table(doc, ['Alert Type', 'Count'], a_rows)

doc.add_paragraph()

# ---- 7. CONCLUSION ----
doc.add_heading('7. Conclusion', level=1)
p = doc.add_paragraph()
p.add_run('SmartStock successfully delivers all four core objectives:')
doc.add_paragraph()
bp(doc, 'Demand Forecasting — Product-level predictions generated using trained ML model')
bp(doc, 'Safety Stock & Reorder Points — Dynamically calculated for each product-brand combination')
bp(doc, 'Slow-Moving & Deadstock — Identified using turnover ratios and sales thresholds')
bp(doc, 'Daily Recommendations — Automated alerts with priority levels and specific actions')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('The interactive dashboard visualizes all insights across 9 pages, providing inventory managers with a comprehensive, data-driven decision support system.')

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Document Version: 1.0  |  Date: April 2025  |  SmartStock — CRISP-ML(Q) Phase 5')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(128,128,128)

doc.save('05_Evaluation_Documentation.docx')
print("✅ Created: 05_Evaluation_Documentation.docx")
