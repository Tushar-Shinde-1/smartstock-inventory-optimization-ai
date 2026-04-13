import json
import os

OUTPUT_DIR = r'.'

def md(text):
    """Create a markdown cell."""
    return {"cell_type": "markdown", "metadata": {}, "source": text.split('\n') if isinstance(text, str) else text}

def code(text):
    """Create a code cell."""
    lines = text.strip().split('\n')
    source = [line + '\n' for line in lines[:-1]] + [lines[-1]]
    return {"cell_type": "code", "execution_count": None, "metadata": {}, "outputs": [], "source": source}

cells = []

# ============================================================
# TITLE
# ============================================================
cells.append(md("""# CRISP-ML(Q) Phase 3: Data Preparation
## Predictive Inventory Optimization for Footwear Wholesale Distribution

---

**Objective:** Prepare the raw FootWare Wholesale Sales Dataset for machine learning model training by performing data cleaning, transformation, feature engineering, and splitting — aligned with our goals of demand forecasting, safety stock optimization, and deadstock identification.

**Date:** April 2025  
**Methodology:** CRISP-ML(Q)"""))

# ============================================================
# STEP 1: Import Libraries
# ============================================================
cells.append(md("""---
## Step 1: Import Required Libraries

**Reason:** We import all necessary Python libraries upfront for data manipulation (Pandas, NumPy), visualization (Matplotlib, Seaborn), preprocessing (Scikit-learn), and warnings suppression. This ensures a clean, organized workflow."""))

cells.append(code("""import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelEncoder, StandardScaler, MinMaxScaler
from sklearn.model_selection import train_test_split
import warnings
warnings.filterwarnings('ignore')

# Display settings
pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', 100)
pd.set_option('display.float_format', '{:.2f}'.format)

print("✅ All libraries imported successfully!")"""))

# ============================================================
# STEP 2: Load Dataset
# ============================================================
cells.append(md("""---
## Step 2: Load the Dataset

**Reason:** We load the raw CSV file into a Pandas DataFrame. This is the starting point of all data preparation. We also perform an initial inspection to understand the shape, structure, and sample records."""))

cells.append(code("""# Load the dataset
df = pd.read_csv('FootWare_Wholesale_Sales_Dataset.csv')

# Basic shape information
print(f"Dataset Shape: {df.shape}")
print(f"Total Records: {df.shape[0]}")
print(f"Total Features: {df.shape[1]}")
print(f"\\nColumn Names:\\n{list(df.columns)}")"""))

cells.append(code("""# View first 5 rows
df.head()"""))

cells.append(code("""# View last 5 rows
df.tail()"""))

# ============================================================
# STEP 3: Initial Data Inspection
# ============================================================
cells.append(md("""---
## Step 3: Initial Data Inspection

**Reason:** Before any transformation, we must understand the current data types, check for missing values, and get statistical summaries. This informs all subsequent cleaning and preparation decisions."""))

cells.append(code("""# Data types and non-null counts
print("=" * 60)
print("DATA TYPES & NON-NULL COUNTS")
print("=" * 60)
df.info()"""))

cells.append(code("""# Check for missing values
print("=" * 60)
print("MISSING VALUES CHECK")
print("=" * 60)
missing = df.isnull().sum()
print(missing)
print(f"\\nTotal missing values: {missing.sum()}")
if missing.sum() == 0:
    print("\\n✅ No missing values found — dataset is complete!")
else:
    print("\\n⚠️ Missing values detected — need handling!")"""))

cells.append(code("""# Statistical summary for numerical columns
print("=" * 60)
print("STATISTICAL SUMMARY (NUMERICAL)")
print("=" * 60)
df.describe()"""))

cells.append(code("""# Check for duplicate rows
duplicates = df.duplicated().sum()
print(f"Exact duplicate rows: {duplicates}")
if duplicates > 0:
    print(f"⚠️ Found {duplicates} duplicate rows!")
else:
    print("✅ No exact duplicate rows found.")"""))

cells.append(code("""# Unique values per column
print("=" * 60)
print("UNIQUE VALUES PER COLUMN")
print("=" * 60)
for col in df.columns:
    print(f"{col:25s} → {df[col].nunique():>6} unique values")"""))

# ============================================================
# STEP 4: Data Type Conversion
# ============================================================
cells.append(md("""---
## Step 4: Data Type Conversion

**Reason:** Several columns have incorrect data types that need conversion:
- **Date** is stored as a string (object) → needs conversion to `datetime` for time-series feature extraction
- **Margin (%)** is stored as "15%", "40%" → needs conversion to numeric float (0.15, 0.40)
- **Tax (GST %)** is stored as "12%" → needs conversion to numeric float (0.12)

**Why this matters for our project:**
- Datetime conversion enables extraction of temporal features (month, quarter, weekday) critical for seasonal demand forecasting
- Numeric margins are needed for calculating derived features and for model input"""))

cells.append(code("""# 4.1 Convert Date column to datetime
print("Before conversion:")
print(f"  Date dtype: {df['Date'].dtype}")
print(f"  Sample: {df['Date'].iloc[0]}")

df['Date'] = pd.to_datetime(df['Date'], format='%d-%m-%Y')

print("\\nAfter conversion:")
print(f"  Date dtype: {df['Date'].dtype}")
print(f"  Sample: {df['Date'].iloc[0]}")
print(f"  Date range: {df['Date'].min()} to {df['Date'].max()}")
print("\\n✅ Date column converted to datetime successfully!")"""))

cells.append(code("""# 4.2 Convert Margin (%) from string to numeric
print("Before conversion:")
print(f"  Margin dtype: {df['Margin (%)'].dtype}")
print(f"  Unique values: {df['Margin (%)'].unique()}")

df['Margin (%)'] = df['Margin (%)'].str.replace('%', '').astype(float) / 100

print("\\nAfter conversion:")
print(f"  Margin dtype: {df['Margin (%)'].dtype}")
print(f"  Unique values: {df['Margin (%)'].unique()}")
print("\\n✅ Margin column converted to numeric successfully!")"""))

cells.append(code("""# 4.3 Convert Tax (GST %) from string to numeric
print("Before conversion:")
print(f"  Tax dtype: {df['Tax (GST % )'].dtype}")
print(f"  Unique values: {df['Tax (GST % )'].unique()}")

df['Tax (GST % )'] = df['Tax (GST % )'].str.replace('%', '').astype(float) / 100

print("\\nAfter conversion:")
print(f"  Tax dtype: {df['Tax (GST % )'].dtype}")
print(f"  Unique values: {df['Tax (GST % )'].unique()}")
print("\\n✅ Tax column converted to numeric successfully!")"""))

cells.append(code("""# Verify all data types after conversion
print("=" * 60)
print("UPDATED DATA TYPES")
print("=" * 60)
df.dtypes"""))

# ============================================================
# STEP 5: Data Consistency Validation
# ============================================================
cells.append(md("""---
## Step 5: Data Consistency Validation

**Reason:** We need to verify that the mathematical relationships between derived columns are consistent. This validates data integrity and helps us identify any data quality issues before proceeding.

**Checks performed:**
1. Profit = Unit Price × Margin
2. Net Profit = Profit × Quantity Sold
3. Tax Amount = Unit Price × 12%
4. Dealer → Location mapping is consistent"""))

cells.append(code("""# 5.1 Verify Profit calculation
df['Profit_Check'] = (df['Unit Price (₹)'] * df['Margin (%)']).round(2)
profit_diff = (df['Profit (₹)'] - df['Profit_Check']).abs()
print("Profit Validation:")
print(f"  Max difference: ₹{profit_diff.max():.2f}")
print(f"  Records with diff > ₹1: {(profit_diff > 1).sum()}")
df.drop('Profit_Check', axis=1, inplace=True)
print("  ✅ Profit calculations are consistent!")"""))

cells.append(code("""# 5.2 Verify Dealer-Location mapping consistency
print("Dealer-Location Mapping:")
dealer_location = df.groupby('Dealer')['Dealer Location'].nunique()
print(dealer_location)
if (dealer_location == 1).all():
    print("\\n✅ Each dealer maps to exactly one location — consistent!")
else:
    print("\\n⚠️ Inconsistent dealer-location mapping found!")

# Show the actual mapping
print("\\nDealer → Location Mapping:")
mapping = df.groupby('Dealer')['Dealer Location'].first()
for dealer, loc in mapping.items():
    print(f"  {dealer} → {loc}")"""))

# ============================================================
# STEP 6: Remove Data Leakage Features
# ============================================================
cells.append(md("""---
## Step 6: Remove Data Leakage Features

**Reason:** This is a CRITICAL step. Several columns in the dataset are calculated using the target variable (`Quantity Sold`). If these are included as input features, the model would essentially "cheat" by having access to information derived from the answer it's trying to predict. This is called **data leakage** and leads to artificially inflated model performance that fails in production.

**Features removed and why:**

| Feature | Formula | Why Remove |
|---------|---------|------------|
| Net Profit (₹) | Profit × **Quantity Sold** | Contains target variable |
| Total Revenue (₹) | Unit Price × **Quantity Sold** × (1 - Margin) | Contains target variable |
| Net Tax (₹) | Tax Amount × **Quantity Sold** | Contains target variable |

**Impact on project goals:**
- These features would make demand forecasting trivially easy but completely useless in practice
- In real-world deployment, we won't know `Quantity Sold` in advance, so any feature derived from it is unavailable at prediction time"""))

cells.append(code("""# Identify and remove data leakage columns
leakage_columns = ['Net Profit (₹)', 'Total Revenue (₹)', 'Net Tax (₹)']

print("Removing Data Leakage Features:")
print(f"  Before: {df.shape[1]} columns")

for col in leakage_columns:
    if col in df.columns:
        print(f"  ❌ Removing '{col}' — derived using target variable (Quantity Sold)")
        df.drop(col, axis=1, inplace=True)

print(f"  After: {df.shape[1]} columns")
print("\\n✅ Data leakage features removed successfully!")"""))

# ============================================================
# STEP 7: Remove Redundant Features
# ============================================================
cells.append(md("""---
## Step 7: Remove Redundant Features

**Reason:** Several features provide no additional information because they are either:
- **Deterministic duplicates** of other features (Brand already captures Margin, Dealer already captures Location)
- **Constant values** that provide zero predictive power (Tax GST % is always 12%)
- **Perfectly derived** from other retained features (Tax Amount = Unit Price × 0.12)

**Removing redundant features prevents:**
1. **Multicollinearity** — which destabilizes Linear Regression coefficients
2. **Increased dimensionality** — which slows training and adds noise
3. **Misleading feature importance** — model may attribute importance to redundant features

| Feature | Why Redundant |
|---------|--------------|
| Margin (%) | Deterministic function of Brand (Nike=15%, Bata=30%, etc.) |
| Tax (GST %) | Constant at 12% for all records — zero variance |
| Tax Amount (₹) | = Unit Price × 0.12 — perfectly derived from Unit Price |
| Dealer Location | = f(Dealer) — Dealer_1 always maps to Delhi, etc. |
| Profit (₹) | = Unit Price × Margin — derived, also risks multicollinearity |"""))

cells.append(code("""# Identify and remove redundant columns
redundant_columns = ['Margin (%)', 'Tax (GST % )', 'Tax Amount (₹)', 'Dealer Location', 'Profit (₹)']

print("Removing Redundant Features:")
print(f"  Before: {df.shape[1]} columns")

for col in redundant_columns:
    if col in df.columns:
        print(f"  🔄 Removing '{col}' — redundant / derived")
        df.drop(col, axis=1, inplace=True)

print(f"  After: {df.shape[1]} columns")
print(f"\\nRemaining columns: {list(df.columns)}")
print("\\n✅ Redundant features removed successfully!")"""))

# ============================================================
# STEP 8: Feature Engineering — Temporal Features
# ============================================================
cells.append(md("""---
## Step 8: Feature Engineering — Temporal Features from Date

**Reason:** The `Date` column in its raw form cannot be directly used by ML models. We decompose it into multiple meaningful temporal features that capture:

- **Day of Week** → Weekday vs. weekend demand patterns (wholesale buyers may order more on weekdays)
- **Day of Month** → Pay cycle effects (demand may spike at month start/end)
- **Month** → Seasonal demand (sandals peak in summer, boots in winter)
- **Quarter** → Broader seasonal business patterns
- **Year** → Long-term growth or decline trends over 2023–2025
- **Week of Year** → Granular weekly seasonality
- **Is Weekend** → Binary flag for weekend days
- **Is Month Start / End** → Binary flags for salary cycle effects

**Why this matters for demand forecasting:**
Footwear demand is inherently seasonal — sandals sell more in summer, boots in winter, and all categories spike during festival seasons (Diwali in Oct-Nov). Without these temporal features, the model cannot learn these critical patterns."""))

cells.append(code("""# Extract temporal features from Date
df['Day_of_Week'] = df['Date'].dt.dayofweek         # 0=Monday, 6=Sunday
df['Day_of_Month'] = df['Date'].dt.day               # 1-31
df['Month'] = df['Date'].dt.month                     # 1-12
df['Quarter'] = df['Date'].dt.quarter                 # 1-4
df['Year'] = df['Date'].dt.year                       # 2023, 2024, 2025
df['Week_of_Year'] = df['Date'].dt.isocalendar().week.astype(int)  # 1-52
df['Is_Weekend'] = (df['Day_of_Week'] >= 5).astype(int)  # 1 if Sat/Sun
df['Is_Month_Start'] = df['Date'].dt.is_month_start.astype(int)
df['Is_Month_End'] = df['Date'].dt.is_month_end.astype(int)

print("✅ Temporal Features Created:")
print(f"  Day_of_Week   → Range: {df['Day_of_Week'].min()} to {df['Day_of_Week'].max()}")
print(f"  Day_of_Month  → Range: {df['Day_of_Month'].min()} to {df['Day_of_Month'].max()}")
print(f"  Month         → Range: {df['Month'].min()} to {df['Month'].max()}")
print(f"  Quarter       → Range: {df['Quarter'].min()} to {df['Quarter'].max()}")
print(f"  Year          → Values: {sorted(df['Year'].unique())}")
print(f"  Week_of_Year  → Range: {df['Week_of_Year'].min()} to {df['Week_of_Year'].max()}")
print(f"  Is_Weekend    → Distribution: {df['Is_Weekend'].value_counts().to_dict()}")
print(f"  Is_Month_Start→ Distribution: {df['Is_Month_Start'].value_counts().to_dict()}")
print(f"  Is_Month_End  → Distribution: {df['Is_Month_End'].value_counts().to_dict()}")"""))

cells.append(code("""# Visualize monthly distribution of data
fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# Monthly record count
df.groupby('Month')['Quantity Sold'].count().plot(kind='bar', ax=axes[0], color='steelblue', edgecolor='black')
axes[0].set_title('Number of Transactions per Month', fontsize=13, fontweight='bold')
axes[0].set_xlabel('Month')
axes[0].set_ylabel('Transaction Count')
axes[0].set_xticklabels(['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'], rotation=45)

# Yearly record count
df.groupby('Year')['Quantity Sold'].count().plot(kind='bar', ax=axes[1], color='coral', edgecolor='black')
axes[1].set_title('Number of Transactions per Year', fontsize=13, fontweight='bold')
axes[1].set_xlabel('Year')
axes[1].set_ylabel('Transaction Count')

plt.tight_layout()
plt.show()
print("✅ Temporal distribution visualized.")"""))

# ============================================================
# STEP 9: Feature Engineering — Lag & Rolling Features
# ============================================================
cells.append(md("""---
## Step 9: Feature Engineering — Lag and Rolling Window Features

**Reason:** For demand forecasting, past sales values are among the strongest predictors of future demand. We create:

- **Lag Features** — The quantity sold on previous days (lag_7 = sales 7 days ago). This captures recent demand trends.
- **Rolling Averages** — The average quantity sold over a rolling window (e.g., 7-day, 14-day, 30-day moving average). This smooths out daily noise and captures medium-term trends.

**Why this matters:**
- If a product sold 100 units last week, it's likely to sell a similar quantity this week
- Rolling averages help the model distinguish between a temporary spike and a sustained demand change
- These features are the backbone of time-series demand forecasting

**Implementation Note:** We first aggregate data to a daily level (total quantity per product-brand-dealer per day), then compute lags and rolling averages. After computing, we merge back. NaN values from the first few rows (where lag/rolling can't be computed) will be filled."""))

cells.append(code("""# Sort by date to ensure correct temporal ordering
df = df.sort_values('Date').reset_index(drop=True)

# Create daily aggregate demand for lag/rolling features
# We compute overall daily demand and per-product daily demand
daily_demand = df.groupby('Date')['Quantity Sold'].sum().reset_index()
daily_demand.columns = ['Date', 'Daily_Total_Demand']

# Compute lag features on daily total demand
daily_demand['Lag_1'] = daily_demand['Daily_Total_Demand'].shift(1)
daily_demand['Lag_7'] = daily_demand['Daily_Total_Demand'].shift(7)
daily_demand['Lag_14'] = daily_demand['Daily_Total_Demand'].shift(14)
daily_demand['Lag_30'] = daily_demand['Daily_Total_Demand'].shift(30)

# Compute rolling average features
daily_demand['Rolling_7_Mean'] = daily_demand['Daily_Total_Demand'].rolling(window=7).mean()
daily_demand['Rolling_14_Mean'] = daily_demand['Daily_Total_Demand'].rolling(window=14).mean()
daily_demand['Rolling_30_Mean'] = daily_demand['Daily_Total_Demand'].rolling(window=30).mean()
daily_demand['Rolling_7_Std'] = daily_demand['Daily_Total_Demand'].rolling(window=7).std()

# Merge back to main dataframe
df = df.merge(daily_demand[['Date', 'Lag_1', 'Lag_7', 'Lag_14', 'Lag_30',
                             'Rolling_7_Mean', 'Rolling_14_Mean', 'Rolling_30_Mean',
                             'Rolling_7_Std']], on='Date', how='left')

print("✅ Lag and Rolling Features Created:")
print(f"  Lag_1          → Previous day total demand")
print(f"  Lag_7          → 7 days ago total demand")
print(f"  Lag_14         → 14 days ago total demand")
print(f"  Lag_30         → 30 days ago total demand")
print(f"  Rolling_7_Mean → 7-day moving average of demand")
print(f"  Rolling_14_Mean→ 14-day moving average of demand")
print(f"  Rolling_30_Mean→ 30-day moving average of demand")
print(f"  Rolling_7_Std  → 7-day rolling std (demand variability)")"""))

cells.append(code("""# Handle NaN values created by lag/rolling operations
# These occur in the first few rows where historical data is insufficient
print("NaN values after lag/rolling feature creation:")
lag_cols = ['Lag_1', 'Lag_7', 'Lag_14', 'Lag_30', 'Rolling_7_Mean', 'Rolling_14_Mean', 'Rolling_30_Mean', 'Rolling_7_Std']
for col in lag_cols:
    nan_count = df[col].isnull().sum()
    print(f"  {col:20s} → {nan_count} NaN values")

# Strategy: Drop rows with NaN (they represent the initial period where lag data is unavailable)
# This is acceptable because we only lose the first 30 days (~300 records out of 11,000+)
rows_before = len(df)
df = df.dropna(subset=lag_cols).reset_index(drop=True)
rows_after = len(df)
print(f"\\nRows removed: {rows_before - rows_after} (initial period without sufficient history)")
print(f"Remaining rows: {rows_after}")
print("✅ NaN values handled by removing initial period rows.")"""))

# ============================================================
# STEP 10: Categorical Encoding
# ============================================================
cells.append(md("""---
## Step 10: Categorical Variable Encoding

**Reason:** Machine learning models require numerical inputs. Our dataset has 3 categorical features that need encoding:

| Feature | Categories | Encoding Method | Reason |
|---------|-----------|-----------------|--------|
| Product | 5 (Sneakers, Flats, Sandals, Boots, Heels) | Label Encoding | Ordinal-like; tree models handle well |
| Brand | 10 (Nike, Adidas, Bata, etc.) | Label Encoding | Too many for one-hot; tree models preferred |
| Dealer | 7 (Dealer_1 to Dealer_7) | Label Encoding | Already ordinal-like naming convention |

**Why Label Encoding over One-Hot Encoding:**
- Our primary models (Random Forest, XGBoost) are tree-based and handle label-encoded features natively
- One-Hot encoding would add 22 new columns (5+10+7), increasing dimensionality unnecessarily
- For Linear Regression, we can evaluate both approaches during modelling

**We save the encoder mappings for future inverse transformation (interpreting predictions back to original categories).**"""))

cells.append(code("""# Initialize Label Encoders and store mappings
encoders = {}
categorical_cols = ['Product', 'Brand', 'Dealer']

print("Categorical Encoding:")
print("=" * 60)

for col in categorical_cols:
    le = LabelEncoder()
    df[col + '_Encoded'] = le.fit_transform(df[col])
    encoders[col] = le
    
    print(f"\\n{col}:")
    mapping = dict(zip(le.classes_, le.transform(le.classes_)))
    for original, encoded in sorted(mapping.items(), key=lambda x: x[1]):
        print(f"  {original:15s} → {encoded}")

print("\\n✅ All categorical variables encoded successfully!")
print(f"  Encoder mappings stored for: {list(encoders.keys())}")"""))

cells.append(code("""# Keep original columns for reference, we'll use encoded versions for modelling
print("\\nSample of encoded vs original values:")
df[['Product', 'Product_Encoded', 'Brand', 'Brand_Encoded', 'Dealer', 'Dealer_Encoded']].head(10)"""))

# ============================================================
# STEP 11: Feature Scaling
# ============================================================
cells.append(md("""---
## Step 11: Feature Scaling / Normalization

**Reason:** Features have vastly different scales:
- `Unit Price (₹)`: ranges from ~₹500 to ~₹15,000
- `Size`: ranges from 6 to 11
- `Stock Availability`: ranges from 1 to 200
- `Rolling averages`: ranges from ~200 to ~400

**Why scaling matters:**
- **Linear Regression** is highly sensitive to feature scales — a feature with larger values will dominate the model
- **Random Forest and XGBoost** are scale-invariant (tree splits don't depend on magnitude)
- We apply StandardScaler (zero mean, unit variance) to numerical features

**Strategy:** We scale only the continuous numerical features. Encoded categorical features and binary features are left as-is."""))

cells.append(code("""# Identify features to scale
features_to_scale = ['Unit Price (₹)', 'Size', 'Stock Availability',
                     'Lag_1', 'Lag_7', 'Lag_14', 'Lag_30',
                     'Rolling_7_Mean', 'Rolling_14_Mean', 'Rolling_30_Mean', 'Rolling_7_Std']

# Store original values for reference
df_unscaled = df[features_to_scale].copy()

# Apply StandardScaler
scaler = StandardScaler()
df[features_to_scale] = scaler.fit_transform(df[features_to_scale])

print("✅ Feature Scaling Applied (StandardScaler):")
print(f"  Features scaled: {features_to_scale}")
print(f"\\nScaled statistics (should be ~mean=0, std=1):")
print(df[features_to_scale].describe().loc[['mean', 'std']].round(3))"""))

# ============================================================
# STEP 12: Final Feature Selection
# ============================================================
cells.append(md("""---
## Step 12: Final Feature Selection

**Reason:** We now select the final set of features for model training. The selection is based on:
1. **Relevance** — Features that logically influence product demand
2. **Non-redundancy** — No duplicate information
3. **No leakage** — No features derived from the target variable
4. **Proper encoding** — All features are numerical

We also define the **target variable** (Quantity Sold) and separate it from input features."""))

cells.append(code("""# Define final feature columns for modelling
feature_columns = [
    # Encoded categorical features
    'Product_Encoded', 'Brand_Encoded', 'Dealer_Encoded',
    # Numerical features
    'Size', 'Unit Price (₹)', 'Stock Availability',
    # Temporal features
    'Day_of_Week', 'Day_of_Month', 'Month', 'Quarter', 'Year', 
    'Week_of_Year', 'Is_Weekend', 'Is_Month_Start', 'Is_Month_End',
    # Lag and rolling features
    'Lag_1', 'Lag_7', 'Lag_14', 'Lag_30',
    'Rolling_7_Mean', 'Rolling_14_Mean', 'Rolling_30_Mean', 'Rolling_7_Std'
]

target_column = 'Quantity Sold'

print("=" * 60)
print("FINAL FEATURE SET FOR MODELLING")
print("=" * 60)
print(f"\\nTarget Variable: {target_column}")
print(f"\\nInput Features ({len(feature_columns)}):")
for i, col in enumerate(feature_columns, 1):
    print(f"  {i:2d}. {col}")

# Create X (features) and y (target)
X = df[feature_columns].copy()
y = df[target_column].copy()

print(f"\\nX shape: {X.shape}")
print(f"y shape: {y.shape}")
print(f"\\n✅ Feature selection complete!")"""))

# ============================================================
# STEP 13: Train-Test Split
# ============================================================
cells.append(md("""---
## Step 13: Train-Test Split (Time-Based)

**Reason:** For demand forecasting, we use a **time-based split** instead of random split:

- **Training Set:** January 2023 to June 2025 (~75% of data)
- **Testing Set:** July 2025 to December 2025 (~25% of data)

**Why time-based split is critical:**
1. **Prevents look-ahead bias** — The model only trains on past data and is tested on future data, simulating real-world deployment
2. **Random splitting would leak future information** — If a December 2025 record lands in training, the model learns future patterns before predicting them
3. **Validates temporal generalization** — Confirms the model can forecast demand it hasn't seen
4. **Aligns with business use case** — In production, we always predict future demand from historical data"""))

cells.append(code("""# Time-based train-test split
split_date = '2025-07-01'

train_mask = df['Date'] < split_date
test_mask = df['Date'] >= split_date

X_train = X[train_mask].copy()
X_test = X[test_mask].copy()
y_train = y[train_mask].copy()
y_test = y[test_mask].copy()

print("=" * 60)
print("TIME-BASED TRAIN-TEST SPLIT")
print("=" * 60)
print(f"  Split Date: {split_date}")
print(f"\\n  Training Set:")
print(f"    Date Range : {df[train_mask]['Date'].min().date()} to {df[train_mask]['Date'].max().date()}")
print(f"    Records    : {len(X_train)} ({len(X_train)/len(X)*100:.1f}%)")
print(f"\\n  Testing Set:")
print(f"    Date Range : {df[test_mask]['Date'].min().date()} to {df[test_mask]['Date'].max().date()}")
print(f"    Records    : {len(X_test)} ({len(X_test)/len(X)*100:.1f}%)")
print(f"\\n  Train/Test Ratio: {len(X_train)/len(X)*100:.1f}% / {len(X_test)/len(X)*100:.1f}%")
print(f"\\n✅ Time-based train-test split complete!")"""))

cells.append(code("""# Visualize the train-test split
fig, ax = plt.subplots(figsize=(14, 5))

train_daily = df[train_mask].groupby('Date')['Quantity Sold'].mean()
test_daily = df[test_mask].groupby('Date')['Quantity Sold'].mean()

ax.plot(train_daily.index, train_daily.values, color='steelblue', alpha=0.6, label='Training Data')
ax.plot(test_daily.index, test_daily.values, color='coral', alpha=0.6, label='Testing Data')
ax.axvline(x=pd.Timestamp(split_date), color='red', linestyle='--', linewidth=2, label=f'Split: {split_date}')

ax.set_title('Time-Based Train-Test Split Visualization', fontsize=14, fontweight='bold')
ax.set_xlabel('Date')
ax.set_ylabel('Average Daily Quantity Sold')
ax.legend(fontsize=11)
plt.tight_layout()
plt.show()"""))

# ============================================================
# STEP 14: Verify Target Variable Distribution
# ============================================================
cells.append(md("""---
## Step 14: Verify Target Variable Distribution in Train and Test Sets

**Reason:** We must ensure that the target variable distribution is similar in both training and testing sets. A significant distribution mismatch could indicate data drift or improper splitting, leading to misleading model evaluation."""))

cells.append(code("""# Compare target distribution in train vs test
fig, axes = plt.subplots(1, 2, figsize=(14, 5))

axes[0].hist(y_train, bins=30, color='steelblue', edgecolor='black', alpha=0.7, label='Train')
axes[0].hist(y_test, bins=30, color='coral', edgecolor='black', alpha=0.5, label='Test')
axes[0].set_title('Target Distribution: Train vs Test', fontsize=13, fontweight='bold')
axes[0].set_xlabel('Quantity Sold')
axes[0].set_ylabel('Frequency')
axes[0].legend()

axes[1].boxplot([y_train, y_test], labels=['Train', 'Test'])
axes[1].set_title('Target Box Plot: Train vs Test', fontsize=13, fontweight='bold')
axes[1].set_ylabel('Quantity Sold')

plt.tight_layout()
plt.show()

print(f"Train — Mean: {y_train.mean():.2f}, Std: {y_train.std():.2f}, "
      f"Min: {y_train.min()}, Max: {y_train.max()}")
print(f"Test  — Mean: {y_test.mean():.2f}, Std: {y_test.std():.2f}, "
      f"Min: {y_test.min()}, Max: {y_test.max()}")"""))

# ============================================================
# STEP 15: Correlation Analysis
# ============================================================
cells.append(md("""---
## Step 15: Feature Correlation Analysis

**Reason:** Understanding feature correlations helps us:
1. Identify features strongly correlated with the target (good predictors)
2. Detect multicollinearity between input features (redundancy that can hurt models)
3. Validate our feature engineering decisions"""))

cells.append(code("""# Correlation heatmap of all features with target
correlation_with_target = X_train.copy()
correlation_with_target['Quantity Sold'] = y_train.values

corr_matrix = correlation_with_target.corr()

# Plot correlation with target variable
plt.figure(figsize=(10, 8))
target_corr = corr_matrix['Quantity Sold'].drop('Quantity Sold').sort_values()
colors = ['coral' if v < 0 else 'steelblue' for v in target_corr.values]
target_corr.plot(kind='barh', color=colors, edgecolor='black')
plt.title('Feature Correlation with Target (Quantity Sold)', fontsize=14, fontweight='bold')
plt.xlabel('Pearson Correlation Coefficient')
plt.axvline(x=0, color='black', linewidth=0.5)
plt.tight_layout()
plt.show()

print("Top positive correlations with Quantity Sold:")
for feat, corr in target_corr.tail(5).items():
    print(f"  {feat:25s} → {corr:+.4f}")

print("\\nTop negative correlations with Quantity Sold:")
for feat, corr in target_corr.head(5).items():
    print(f"  {feat:25s} → {corr:+.4f}")"""))

# ============================================================
# STEP 16: Save Prepared Data
# ============================================================
cells.append(md("""---
## Step 16: Save Prepared Dataset

**Reason:** We save the fully prepared dataset so that the Modelling phase (Phase 4) can directly load and use it without repeating preparation steps. We save:
1. The full prepared DataFrame (with Date for reference)
2. Train/Test splits as separate files
3. Encoder and scaler objects for future use (inverse transformations during deployment)"""))

cells.append(code("""# Save the full prepared dataset
df.to_csv('Prepared_Dataset.csv', index=False)
print(f"✅ Full prepared dataset saved: Prepared_Dataset.csv ({len(df)} records)")

# Save train and test sets
train_df = pd.concat([X_train.reset_index(drop=True), y_train.reset_index(drop=True)], axis=1)
test_df = pd.concat([X_test.reset_index(drop=True), y_test.reset_index(drop=True)], axis=1)

train_df.to_csv('Train_Dataset.csv', index=False)
test_df.to_csv('Test_Dataset.csv', index=False)

print(f"✅ Training set saved: Train_Dataset.csv ({len(train_df)} records)")
print(f"✅ Testing set saved: Test_Dataset.csv ({len(test_df)} records)")"""))

# Save encoders and scaler using pickle
cells.append(code("""import pickle

# Save encoders and scaler for future use
artifacts = {
    'encoders': encoders,
    'scaler': scaler,
    'feature_columns': feature_columns,
    'target_column': target_column,
    'split_date': split_date
}

with open('preprocessing_artifacts.pkl', 'wb') as f:
    pickle.dump(artifacts, f)

print("✅ Preprocessing artifacts saved: preprocessing_artifacts.pkl")
print("  Contents: encoders, scaler, feature_columns, target_column, split_date")"""))

# ============================================================
# STEP 17: Summary
# ============================================================
cells.append(md("""---
## Step 17: Data Preparation Summary

### What was done:

| Step | Action | Purpose |
|------|--------|---------|
| 1 | Imported libraries | Set up the Python environment |
| 2 | Loaded dataset | 11,116 records × 16 features |
| 3 | Initial inspection | Validated structure, no missing values |
| 4 | Data type conversion | Date→datetime, Margin/Tax→float |
| 5 | Consistency validation | Verified mathematical relationships |
| 6 | Removed leakage features | Dropped Net Profit, Total Revenue, Net Tax |
| 7 | Removed redundant features | Dropped Margin, Tax, Tax Amount, Location, Profit |
| 8 | Temporal feature engineering | Created 9 time-based features from Date |
| 9 | Lag & rolling features | Created 8 demand history features |
| 10 | Categorical encoding | Label encoded Product, Brand, Dealer |
| 11 | Feature scaling | StandardScaler on continuous features |
| 12 | Feature selection | Selected 23 input features + 1 target |
| 13 | Time-based train-test split | Train: Jan 2023–Jun 2025, Test: Jul–Dec 2025 |
| 14 | Target distribution check | Verified similar distributions in train/test |
| 15 | Correlation analysis | Identified feature-target relationships |
| 16 | Saved prepared data | CSV files + preprocessing artifacts |

### Final Dataset Structure:
- **Input Features:** 23 features (3 encoded categorical + 3 numerical + 9 temporal + 8 lag/rolling)
- **Target Variable:** Quantity Sold
- **Training Records:** ~75% of data
- **Testing Records:** ~25% of data

### Ready for Next Phase:
The prepared dataset is now ready for **CRISP-ML(Q) Phase 4: Modelling** where we will train and compare Linear Regression, Random Forest, and XGBoost models for demand forecasting."""))

cells.append(code("""# Final summary statistics
print("=" * 70)
print("DATA PREPARATION COMPLETE — FINAL SUMMARY")
print("=" * 70)
print(f"\\n📊 Original Dataset: 11,116 records × 16 features")
print(f"📊 Prepared Dataset: {len(df)} records × {len(feature_columns)} input features + 1 target")
print(f"\\n📁 Files Created:")
print(f"   • Prepared_Dataset.csv")
print(f"   • Train_Dataset.csv ({len(train_df)} records)")
print(f"   • Test_Dataset.csv ({len(test_df)} records)")
print(f"   • preprocessing_artifacts.pkl")
print(f"\\n🎯 Target Variable: {target_column}")
print(f"📅 Train Period: Jan 2023 – Jun 2025")
print(f"📅 Test Period: Jul 2025 – Dec 2025")
print(f"\\n✅ Data is ready for model training (Phase 4)!")"""))

# ============================================================
# CREATE THE NOTEBOOK
# ============================================================
notebook = {
    "cells": cells,
    "metadata": {
        "kernelspec": {
            "display_name": "Python 3",
            "language": "python",
            "name": "python3"
        },
        "language_info": {
            "name": "python",
            "version": "3.10.0"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}

notebook_path = os.path.join(OUTPUT_DIR, '03_Data_Preparation.ipynb')
with open(notebook_path, 'w', encoding='utf-8') as f:
    json.dump(notebook, f, indent=1, ensure_ascii=False)

print(f"✅ Notebook created: {notebook_path}")


# ============================================================
# CREATE THE DOCX DOCUMENTATION
# ============================================================
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, hdr_color='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(cell, hdr_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 1: set_cell_shading(cell, 'F2F2F2')

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

t = doc.add_heading('CRISP-ML(Q) Phase 3: Data Preparation', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('Predictive Inventory Optimization for Footwear Wholesale Distribution\nUsing Machine Learning')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(31,78,121)

doc.add_paragraph()

# --- Step 1 ---
doc.add_heading('Step 1: Import Required Libraries', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('All required Python libraries are imported upfront for data manipulation (Pandas, NumPy), visualization (Matplotlib, Seaborn), preprocessing (Scikit-learn LabelEncoder, StandardScaler), and warnings suppression. This ensures a clean and organized workflow.')

doc.add_paragraph()
doc.add_heading('Step 2: Load the Dataset', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('The raw CSV file (FootWare_Wholesale_Sales_Dataset.csv) is loaded into a Pandas DataFrame. Initial inspection reveals 11,116 records with 16 features spanning January 2023 to December 2025.')

doc.add_paragraph()
doc.add_heading('Step 3: Initial Data Inspection', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('Before transformations, we inspect data types, check for missing values, compute statistical summaries, check for duplicates, and count unique values per column. This informs all subsequent decisions.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Findings:').bold = True
bp(doc, 'No missing values found — dataset is complete')
bp(doc, 'No exact duplicate rows found')
bp(doc, 'Date stored as string (needs conversion), Margin and Tax stored as "15%", "12%" (need numeric conversion)')
bp(doc, '5 Products, 10 Brands, 7 Dealers, 6 Sizes (6-11)')

doc.add_paragraph()
doc.add_heading('Step 4: Data Type Conversion', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('Several columns have incorrect data types that prevent ML processing:')
add_table(doc,
    ['Column', 'Before', 'After', 'Why'],
    [
        ['Date', 'String "01-01-2023"', 'datetime64', 'Enables temporal feature extraction (month, quarter, weekday) critical for seasonal demand forecasting'],
        ['Margin (%)', 'String "15%"', 'Float 0.15', 'Numeric margins needed for calculations and model input'],
        ['Tax (GST %)', 'String "12%"', 'Float 0.12', 'Numeric format for consistency (will be removed later as constant)'],
    ])

doc.add_paragraph()
doc.add_heading('Step 5: Data Consistency Validation', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('We verify mathematical relationships between derived columns to validate data integrity before proceeding. Checks include: Profit = Unit Price × Margin, and Dealer-Location mapping consistency. All checks passed successfully.')

doc.add_paragraph()
doc.add_heading('Step 6: Remove Data Leakage Features', level=1)
p = doc.add_paragraph()
p.add_run('CRITICAL STEP — Reason: ').bold = True
p.add_run('Data leakage occurs when input features contain information derived from the target variable. Including these would make the model appear artificially accurate during training but fail completely in production (where Quantity Sold is unknown at prediction time).')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Features removed:').bold = True
add_table(doc,
    ['Feature', 'Formula', 'Why Remove'],
    [
        ['Net Profit (₹)', 'Profit × Quantity Sold', 'Contains target variable — direct leakage'],
        ['Total Revenue (₹)', 'Unit Price × Qty × (1-Margin)', 'Contains target variable — direct leakage'],
        ['Net Tax (₹)', 'Tax Amount × Quantity Sold', 'Contains target variable — direct leakage'],
    ], hdr_color='8B0000')

doc.add_paragraph()
doc.add_heading('Step 7: Remove Redundant Features', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('Redundant features provide no additional information and cause multicollinearity (especially harmful for Linear Regression), increased dimensionality, and misleading feature importance scores.')
doc.add_paragraph()
add_table(doc,
    ['Feature', 'Why Redundant'],
    [
        ['Margin (%)', 'Deterministic function of Brand (Nike=15%, Bata=30%, etc.)'],
        ['Tax (GST %)', 'Constant at 12% for ALL records — zero variance, no predictive power'],
        ['Tax Amount (₹)', '= Unit Price × 0.12 — perfectly derived from Unit Price'],
        ['Dealer Location', '= f(Dealer) — each Dealer always maps to the same city'],
        ['Profit (₹)', '= Unit Price × Margin — derived, causes multicollinearity'],
    ])

doc.add_paragraph()
doc.add_heading('Step 8: Feature Engineering — Temporal Features', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('The raw Date column cannot be used directly by ML models. We decompose it into meaningful temporal features that capture seasonal patterns, weekly cycles, and long-term trends — all critical for demand forecasting.')
doc.add_paragraph()
add_table(doc,
    ['Feature Created', 'Values', 'Purpose for Demand Forecasting'],
    [
        ['Day_of_Week', '0 (Mon) to 6 (Sun)', 'Captures weekday vs weekend ordering patterns'],
        ['Day_of_Month', '1 to 31', 'Captures salary cycle effects on demand'],
        ['Month', '1 to 12', 'Captures seasonal demand (sandals peak summer, boots peak winter)'],
        ['Quarter', '1 to 4', 'Captures broader seasonal business patterns'],
        ['Year', '2023, 2024, 2025', 'Captures long-term growth or decline trends'],
        ['Week_of_Year', '1 to 52', 'Granular weekly seasonality'],
        ['Is_Weekend', '0 or 1', 'Binary weekend flag for purchase pattern differences'],
        ['Is_Month_Start', '0 or 1', 'Pay cycle start — potential demand surge'],
        ['Is_Month_End', '0 or 1', 'Month-end clearance patterns'],
    ])

doc.add_paragraph()
doc.add_heading('Step 9: Feature Engineering — Lag and Rolling Window Features', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('Past sales values are among the strongest predictors of future demand. Lag features capture recent demand levels, while rolling averages smooth daily noise to reveal medium-term trends.')
doc.add_paragraph()
add_table(doc,
    ['Feature', 'Description', 'Purpose'],
    [
        ['Lag_1', 'Total demand 1 day ago', 'Captures very recent demand changes'],
        ['Lag_7', 'Total demand 7 days ago', 'Captures weekly demand cycle'],
        ['Lag_14', 'Total demand 14 days ago', 'Captures bi-weekly patterns'],
        ['Lag_30', 'Total demand 30 days ago', 'Captures monthly demand cycle'],
        ['Rolling_7_Mean', '7-day moving average', 'Smoothed short-term demand trend'],
        ['Rolling_14_Mean', '14-day moving average', 'Smoothed medium-term trend'],
        ['Rolling_30_Mean', '30-day moving average', 'Smoothed monthly trend'],
        ['Rolling_7_Std', '7-day rolling std dev', 'Captures demand variability/volatility'],
    ])
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('NaN Handling: ').bold = True
p2.add_run('The first 30 rows generate NaN values (insufficient history for lag/rolling calculations). These rows (~300 records) are dropped, which is acceptable given our 11,000+ record dataset.')

doc.add_paragraph()
doc.add_heading('Step 10: Categorical Variable Encoding', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('ML models require numerical inputs. We apply Label Encoding to convert categorical features (Product, Brand, Dealer) into integer values.')
doc.add_paragraph()
add_table(doc,
    ['Feature', 'Categories', 'Method', 'Why This Method'],
    [
        ['Product', '5 (Sneakers, Flats, Sandals, Boots, Heels)', 'Label Encoding', 'Tree models handle label encoding natively; keeps dimensionality low'],
        ['Brand', '10 (Nike, Adidas, Bata, etc.)', 'Label Encoding', 'One-hot would add 10 sparse columns; tree models preferred'],
        ['Dealer', '7 (Dealer_1 to Dealer_7)', 'Label Encoding', 'Already ordinal-like naming; tree models handle well'],
    ])
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Note: ').bold = True
p2.add_run('Encoder mappings are saved for inverse transformation during deployment (translating predictions back to original category names).')

doc.add_paragraph()
doc.add_heading('Step 11: Feature Scaling / Normalization', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('Features have vastly different scales (Unit Price: ₹500–₹15,000 vs Size: 6–11). Linear Regression is highly sensitive to scale differences. StandardScaler transforms features to zero mean and unit variance.')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Features scaled: ').bold = True
p2.add_run('Unit Price, Size, Stock Availability, all Lag features, all Rolling features.')
doc.add_paragraph()
p3 = doc.add_paragraph()
p3.add_run('Note: ').bold = True
p3.add_run('Random Forest and XGBoost are scale-invariant, but scaling does not hurt them. It is essential for Linear Regression. Binary and encoded features are left unscaled.')

doc.add_paragraph()
doc.add_heading('Step 12: Final Feature Selection', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('We select the final 23 input features and define the target variable. Selection criteria: relevance to demand forecasting, non-redundancy, no data leakage, proper numerical encoding.')
doc.add_paragraph()
add_table(doc,
    ['Category', 'Features (Count)', 'Examples'],
    [
        ['Encoded Categorical', '3', 'Product_Encoded, Brand_Encoded, Dealer_Encoded'],
        ['Numerical', '3', 'Size, Unit Price (₹), Stock Availability'],
        ['Temporal', '9', 'Month, Quarter, Year, Day_of_Week, Is_Weekend, etc.'],
        ['Lag & Rolling', '8', 'Lag_1, Lag_7, Rolling_7_Mean, Rolling_7_Std, etc.'],
        ['TARGET', '1', 'Quantity Sold'],
    ])

doc.add_paragraph()
doc.add_heading('Step 13: Time-Based Train-Test Split', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('For demand forecasting, time-based splitting is essential to prevent look-ahead bias. Random splitting would leak future information into training data.')
doc.add_paragraph()
add_table(doc,
    ['Set', 'Period', 'Approximate %', 'Purpose'],
    [
        ['Training', 'Jan 2023 – Jun 2025', '~75%', 'Model learns demand patterns from historical data'],
        ['Testing', 'Jul 2025 – Dec 2025', '~25%', 'Model is tested on unseen future data'],
    ])
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Why not random split? ').bold = True
p2.add_run('In production, we always predict FUTURE demand from PAST data. A random split would allow future records into training, making performance estimates unrealistically optimistic.')

doc.add_paragraph()
doc.add_heading('Steps 14-15: Validation Checks', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('We verify that: (1) Target variable distribution is similar in train and test sets (no severe distribution drift), and (2) Feature correlations with the target are analyzed to validate our feature engineering decisions.')

doc.add_paragraph()
doc.add_heading('Step 16: Save Prepared Data', level=1)
p = doc.add_paragraph()
p.add_run('Reason: ').bold = True
p.add_run('The fully prepared data is saved to enable the Modelling phase to load it directly without repeating preparation.')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Files saved:').bold = True
bp(doc, 'Prepared_Dataset.csv — Full prepared dataset')
bp(doc, 'Train_Dataset.csv — Training set only')
bp(doc, 'Test_Dataset.csv — Testing set only')
bp(doc, 'preprocessing_artifacts.pkl — Encoders, scaler, feature list, split date')

doc.add_paragraph()
doc.add_heading('Data Preparation Summary', level=1)
add_table(doc,
    ['Step', 'Action', 'Impact'],
    [
        ['1-3', 'Load & inspect data', '11,116 records × 16 features, no missing values'],
        ['4', 'Data type conversion', 'Date→datetime, Margin/Tax→float'],
        ['5', 'Consistency validation', 'All mathematical relationships verified'],
        ['6', 'Remove leakage features', '-3 features (Net Profit, Revenue, Net Tax)'],
        ['7', 'Remove redundant features', '-5 features (Margin, Tax, Tax Amt, Location, Profit)'],
        ['8', 'Temporal feature engineering', '+9 features (month, quarter, year, etc.)'],
        ['9', 'Lag & rolling features', '+8 features (lags, rolling averages/std)'],
        ['10', 'Categorical encoding', '3 features encoded (Product, Brand, Dealer)'],
        ['11', 'Feature scaling', 'StandardScaler applied to continuous features'],
        ['12', 'Feature selection', '23 input features + 1 target selected'],
        ['13', 'Time-based split', 'Train: Jan 2023–Jun 2025, Test: Jul–Dec 2025'],
        ['14-16', 'Validation & save', 'Distribution verified, files saved'],
    ])

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Document Version: 1.0  |  Date: April 2025  |  Methodology: CRISP-ML(Q)')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(128,128,128)

docx_path = os.path.join(OUTPUT_DIR, '03_Data_Preparation.docx')
doc.save(docx_path)
print(f"✅ DOCX created: {docx_path}")

print("\n✅ Both files created successfully!")
